import io
import logging
import os
import tempfile

import PyPDF2
import docx
from fastapi import HTTPException, UploadFile
from markitdown import MarkItDown
from openai import AsyncOpenAI
from sqlalchemy.ext.asyncio import AsyncSession

from models.document import Document
from services.embedding_service import add_chunks_to_store
from utils.text_splitter import merge_small_chunks, split_by_semantic_boundaries, split_by_structure, split_document

logger = logging.getLogger(__name__)
markitdown_client = MarkItDown(enable_plugins=False)


def _extract_with_markitdown(content: bytes, filename: str) -> str:
    suffix = os.path.splitext(filename)[1] or ".bin"
    temp_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name
        result = markitdown_client.convert(temp_path)
        return str(getattr(result, "text_content", "") or "").strip()
    finally:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)


async def parse_document(file: UploadFile) -> str:
    content = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        try:
            text = _extract_with_markitdown(content, file.filename or "document.pdf")
            if text:
                return text
        except Exception:
            logger.warning("MarkItDown PDF parse failed for %s, falling back to PyPDF2", file.filename, exc_info=True)
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if filename.endswith(".docx"):
        try:
            text = _extract_with_markitdown(content, file.filename or "document.docx")
            if text:
                return text
        except Exception:
            logger.warning("MarkItDown DOCX parse failed for %s, falling back to python-docx", file.filename, exc_info=True)
        document = docx.Document(io.BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
    return content.decode("utf-8", errors="ignore").strip()


async def upload_document(file: UploadFile, description: str, client: AsyncOpenAI, db: AsyncSession) -> Document:
    ext = file.filename.rsplit(".", 1)[-1].lower() if file.filename and "." in file.filename else "txt"
    doc = Document(filename=file.filename or "untitled", doc_type=ext, description=description, status="processing")
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    logger.info("Document upload started: id=%s filename=%s type=%s", doc.id, doc.filename, doc.doc_type)

    try:
        text = await parse_document(file)
        if not text:
            raise HTTPException(status_code=400, detail="Document contains no extractable text")
        structure_chunks = split_by_structure(text, ext)
        semantic_chunks = []
        for chunk in structure_chunks:
            semantic_chunks.extend(split_by_semantic_boundaries(chunk))
        merged_chunks = merge_small_chunks(semantic_chunks)
        chunks = split_document(text, ext)
        if not chunks:
            raise HTTPException(status_code=400, detail="Document could not be split into chunks")
        logger.info(
            "Document split summary: id=%s filename=%s structure_chunks=%s semantic_chunks=%s merged_chunks=%s final_chunks=%s",
            doc.id,
            doc.filename,
            len(structure_chunks),
            len(semantic_chunks),
            len(merged_chunks),
            len(chunks),
        )
        await add_chunks_to_store(doc.id, doc.filename, chunks, client)
        doc.chunk_count = len(chunks)
        doc.status = "ready"
        doc.error_message = None
        logger.info("Document upload finished: id=%s filename=%s chunks=%s status=ready", doc.id, doc.filename, doc.chunk_count)
    except Exception as exc:
        doc.status = "error"
        doc.error_message = str(exc)
        logger.exception("Document upload failed: id=%s filename=%s", doc.id, doc.filename)

    await db.commit()
    await db.refresh(doc)
    return doc
